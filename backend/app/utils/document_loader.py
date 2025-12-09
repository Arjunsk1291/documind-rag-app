import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from llama_index.core import Document
from typing import List
import logging

logger = logging.getLogger(__name__)

class DocumentLoader:
    """Utility class to load different document types"""
    
    @staticmethod
    def load_pdf(file_path: str) -> List[Document]:
        """Load PDF document using PyMuPDF"""
        try:
            logger.info(f"Loading PDF: {file_path}")
            
            # Open PDF
            pdf_document = fitz.open(file_path)
            
            # Extract text from all pages
            text_content = []
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text = page.get_text()
                if text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
            
            pdf_document.close()
            
            # Combine all text
            full_text = "\n\n".join(text_content)
            
            logger.info(f"Extracted {len(full_text)} characters from {len(text_content)} pages")
            
            if not full_text.strip():
                raise ValueError("No text could be extracted from PDF")
            
            # Create LlamaIndex Document
            documents = [Document(
                text=full_text,
                metadata={
                    "file_path": file_path,
                    "file_name": os.path.basename(file_path),
                    "file_type": "pdf",
                    "num_pages": len(text_content)
                }
            )]
            
            logger.info(f"Successfully loaded PDF: {file_path}")
            return documents
            
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def load_docx(file_path: str) -> List[Document]:
        """Load DOCX document"""
        try:
            logger.info(f"Loading DOCX: {file_path}")
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            logger.info(f"Extracted {len(text)} characters from DOCX")
            
            if not text.strip():
                raise ValueError("No text could be extracted from DOCX")
            
            documents = [Document(
                text=text,
                metadata={
                    "file_path": file_path,
                    "file_name": os.path.basename(file_path),
                    "file_type": "docx"
                }
            )]
            
            logger.info(f"Successfully loaded DOCX: {file_path}")
            return documents
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def load_txt(file_path: str) -> List[Document]:
        """Load TXT document"""
        try:
            logger.info(f"Loading TXT: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            logger.info(f"Extracted {len(text)} characters from TXT")
            
            if not text.strip():
                raise ValueError("No text could be extracted from TXT file")
            
            documents = [Document(
                text=text,
                metadata={
                    "file_path": file_path,
                    "file_name": os.path.basename(file_path),
                    "file_type": "txt"
                }
            )]
            
            logger.info(f"Successfully loaded TXT: {file_path}")
            return documents
            
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def load_document(file_path: str) -> List[Document]:
        """Load document based on file extension"""
        ext = os.path.splitext(file_path)[1].lower()
        
        loaders = {
            '.pdf': DocumentLoader.load_pdf,
            '.docx': DocumentLoader.load_docx,
            '.txt': DocumentLoader.load_txt,
            '.md': DocumentLoader.load_txt,
        }
        
        loader = loaders.get(ext)
        if not loader:
            raise ValueError(f"Unsupported file type: {ext}")
        
        return loader(file_path)
